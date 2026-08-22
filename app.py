import streamlit as st
import os
import json
import tempfile
from typing import List, Dict, Any
import fitz  # PyMuPDF
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import openai
from docx import Document
import io

# ---------- DeepSeek API Configuration ----------
DEEPSEEK_API_KEY = st.secrets.get("DEEPSEEK_API_KEY", os.getenv("DEEPSEEK_API_KEY", ""))

# नए OpenAI library (>=1.0.0) के लिए
from openai import OpenAI

client = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com"
)

# ---------- Helper Functions ----------
def extract_text_from_pdf(pdf_path):
    """PDF से text निकालें। पहले direct text try करें, अगर न मिले तो OCR करें।"""
    # Direct text extraction
    doc = fitz.open(pdf_path)
    text = "\n".join(page.get_text("text") for page in doc)
    doc.close()

    if text.strip():
        return text.strip(), False

    # OCR fallback
    images = convert_from_path(pdf_path, dpi=300)
    ocr_text = ""
    for img in images:
        ocr_text += pytesseract.image_to_string(img, lang="hin+eng") + "\n"
    return ocr_text.strip(), True

def parse_deepseek_json_response(response_text):
    """DeepSeek से आए text में से JSON array निकालें।"""
    try:
        # पहले सीधा json.loads try करें
        return json.loads(response_text)
    except:
        pass

    # अगर JSON extra text में wrapped है तो निकालें
    import re
    match = re.search(r'\[.*\]', response_text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except:
            pass
    return None

def extract_questions_with_deepseek(text, question_numbers):
    """DeepSeek AI से text में से specified question numbers निकालें।"""
    if not DEEPSEEK_API_KEY:
        st.error("DeepSeek API key नहीं मिला। कृपया Secrets में DEEPSEEK_API_KEY set करें।")
        return []

    prompt = f"""
आप एक शिक्षक सहायक हैं। नीचे दिए गए text में कई प्रश्न हैं। हर प्रश्न में प्रश्न संख्या, प्रश्न text, चार options (A, B, C, D), सही उत्तर (A/B/C/D) और व्याख्या हो सकती है। Text OCR से आया है, इसलिए कुछ अक्षर गलत हो सकते हैं, कृपया उन्हें सुधारें।

आपको निम्न प्रश्न संख्याएँ निकालनी हैं: {question_numbers}

हर प्रश्न का JSON object इस format में दें:
{{
  "expected_number": int,
  "found_number": int,
  "question_text": str,
  "options": {{"A": str, "B": str, "C": str, "D": str}},
  "answer": "A/B/C/D",
  "explanation": str,
  "match": true/false,
  "confidence": "high/medium/low"
}}

अगर कोई expected_number text में नहीं मिलती, तो उसके लिए object में "found_number" को null और "match" को false रखें।  
केवल JSON array return करें, कोई अन्य text नहीं।

Text:
{text}
"""
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that extracts structured data from text."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=4000
        )
        result_text = response.choices[0].message.content
        return parse_deepseek_json_response(result_text) or []
    except Exception as e:
        st.error(f"DeepSeek API error: {e}")
        return []

def fallback_extract_question(text, used_numbers):
    """अगर कोई question number न मिले तो कोई भी unused question निकालें।"""
    if not DEEPSEEK_API_KEY:
        return None
    prompt = f"""
आप एक शिक्षक सहायक हैं। नीचे दिए गए text में कई प्रश्न हैं। कृपया कोई एक प्रश्न निकालें जो नीचे दिए गए used numbers में नहीं है।

Used numbers: {used_numbers}

हर प्रश्न का JSON object इस format में दें:
{{
  "expected_number": null,
  "found_number": int,
  "question_text": str,
  "options": {{"A": str, "B": str, "C": str, "D": str}},
  "answer": "A/B/C/D",
  "explanation": str,
  "match": false,
  "confidence": "medium"
}}

केवल JSON object return करें, कोई array नहीं।

Text:
{text}
"""
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "You extract one question from text."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=1500
        )
        result_text = response.choices[0].message.content
        # JSON object parse करें
        import re
        match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except:
                return None
        return None
    except:
        return None

def create_model_paper_docx(questions):
    """Model paper और Answer Key Word document बनाएं।"""
    doc = Document()
    doc.add_heading('मॉडल पेपर', level=0)

    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"प्रश्न {i}: {q.get('question_text', '')}")
        for letter in ['A', 'B', 'C', 'D']:
            if letter in q.get('options', {}):
                doc.add_paragraph(f"{letter}) {q['options'][letter]}")
        doc.add_paragraph("")

    doc.add_page_break()
    doc.add_heading('उत्तर कुंजी', level=0)
    for i, q in enumerate(questions, 1):
        answer = q.get('answer') or 'N/A'
        explanation = q.get('explanation') or 'N/A'
        doc.add_paragraph(f"प्रश्न {i}: उत्तर ({answer})")
        doc.add_paragraph(f"व्याख्या: {explanation}")
        doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------- Streamlit UI ----------
st.set_page_config(page_title="PDF से मॉडल पेपर जेनरेटर", layout="wide")
st.title("📄 PDF से मॉडल पेपर जेनरेटर (DeepSeek AI + OCR)")
st.markdown("अपनी PDFs upload करें, प्रश्न संख्याएँ चुनें, और DeepSeek AI आपके लिए मॉडल पेपर तैयार कर देगा।")

uploaded_files = st.file_uploader("PDF files चुनें (आप एक साथ कई upload कर सकते हैं)", type="pdf", accept_multiple_files=True)

# Global question numbers
question_numbers_input = st.text_input(
    "प्रत्येक PDF से कौन-सी प्रश्न संख्याएँ निकालनी हैं? (comma separated, जैसे: 1,2 या 3,4)",
    value="1,2"
)

col1, col2 = st.columns(2)
with col1:
    force_ocr = st.checkbox("हमेशा OCR use करें (अगर text garbled है)", value=True)
with col2:
    total_questions = st.empty()

if st.button("🚀 मॉडल पेपर बनाएं", type="primary"):
    if not uploaded_files:
        st.warning("कृपया कम से कम एक PDF upload करें।")
    elif not DEEPSEEK_API_KEY:
        st.error("DeepSeek API key नहीं मिला। कृपया Secrets में DEEPSEEK_API_KEY set करें।")
    else:
        try:
            question_numbers = [int(x.strip()) for x in question_numbers_input.split(',') if x.strip()]
        except:
            st.error("प्रश्न संख्याएँ सही format में नहीं हैं।")
            question_numbers = []

        if not question_numbers:
            st.stop()

        all_questions = []
        progress_bar = st.progress(0)
        status_text = st.empty()

        # अस्थायी फोल्डर में PDFs save करें
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_paths = []
            for uploaded_file in uploaded_files:
                path = os.path.join(tmpdir, uploaded_file.name)
                with open(path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                pdf_paths.append(path)

            total_pdfs = len(pdf_paths)
            for idx, pdf_path in enumerate(pdf_paths):
                status_text.text(f"📄 Processing: {os.path.basename(pdf_path)}")
                progress_bar.progress((idx) / total_pdfs)

                # Text extraction
                if force_ocr:
                    # हमेशा OCR
                    images = convert_from_path(pdf_path, dpi=300)
                    text = ""
                    for img in images:
                        text += pytesseract.image_to_string(img, lang="hin+eng") + "\n"
                else:
                    text, _ = extract_text_from_pdf(pdf_path)

                # DeepSeek से extraction
                extracted = extract_questions_with_deepseek(text, question_numbers)

                # अगर कुछ expected numbers missing हों तो fallback से भरें
                if len(extracted) < len(question_numbers):
                    used_numbers = {q.get('found_number') for q in extracted if q.get('found_number') is not None}
                    missing_count = len(question_numbers) - len(extracted)
                    for _ in range(missing_count):
                        fallback_q = fallback_extract_question(text, used_numbers)
                        if fallback_q:
                            extracted.append(fallback_q)
                            if fallback_q.get('found_number'):
                                used_numbers.add(fallback_q['found_number'])

                # सिर्फ valid questions add करें
                for q in extracted:
                    if q.get('question_text') and q.get('found_number') is not None:
                        q['source_pdf'] = os.path.basename(pdf_path)
                        all_questions.append(q)

                progress_bar.progress((idx + 1) / total_pdfs)

        status_text.text("✅ सभी PDFs process हो गए।")
        progress_bar.empty()

        if all_questions:
            st.success(f"कुल {len(all_questions)} प्रश्न निकाले गए।")

            # Model paper Word file बनाएं
            docx_buffer = create_model_paper_docx(all_questions)

            st.download_button(
                label="📥 मॉडल पेपर डाउनलोड करें (Word)",
                data=docx_buffer,
                file_name="model_paper.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # JSON download
            json_str = json.dumps(all_questions, ensure_ascii=False, indent=2)
            st.download_button(
                label="📥 JSON डेटा डाउनलोड करें",
                data=json_str,
                file_name="extracted_questions.json",
                mime="application/json"
            )

            # Preview
            with st.expander("🔍 प्रश्न Preview देखें"):
                for i, q in enumerate(all_questions, 1):
                    st.markdown(f"**प्रश्न {i}:** {q['question_text']}")
                    for letter in ['A', 'B', 'C', 'D']:
                        if letter in q.get('options', {}):
                            st.write(f"{letter}) {q['options'][letter]}")
                    st.write(f"उत्तर: {q.get('answer')} | व्याख्या: {q.get('explanation', 'N/A')}")
                    st.divider()
        else:
            st.error("कोई प्रश्न नहीं निकाला जा सका।")
